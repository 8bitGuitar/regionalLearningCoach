import gradio as gr
from pinecone import Pinecone, ServerlessSpec
import os
import tempfile
import shutil
import time
from typing import List, Optional, Tuple
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
import numpy as np
from dataclasses import dataclass
import nltk
from nltk.tokenize import sent_tokenize
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

@dataclass
class Document:
    text: str
    source: str

@dataclass
class TextChunk:
    text: str
    source: str
    embedding: Optional[np.ndarray] = None

class DocumentProcessor:
    def __init__(self):
        self.model = SentenceTransformer('AkshitaS/bhasha-embed-v0')
        self.embedding_dim = 768
    
    def read_file(self, file_path: str) -> Document:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_ext == '.pdf':
            text = ''
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'
        elif file_ext == '.docx':
            doc = docx.Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
            
        return Document(text=text, source=file_path)

    def chunk_text(self, document: Document, chunk_size: int = 768, chunk_overlap: int = 20) -> List[TextChunk]:
        try:
            sentences = sent_tokenize(document.text)
            chunks = []
            current_chunk = []
            current_length = 0
            
            for sentence in sentences:
                sentence_length = len(sentence)
                
                if current_length + sentence_length > chunk_size and current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    chunks.append(TextChunk(
                        text=chunk_text,
                        source=document.source
                    ))
                    
                    overlap_text_length = 0
                    overlap_sentences = []
                    for sent in reversed(current_chunk):
                        if overlap_text_length + len(sent) > chunk_overlap:
                            break
                        overlap_sentences.insert(0, sent)
                        overlap_text_length += len(sent)
                    
                    current_chunk = overlap_sentences
                    current_length = overlap_text_length
                
                current_chunk.append(sentence)
                current_length += sentence_length
            
            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append(TextChunk(
                    text=chunk_text,
                    source=document.source
                ))
            
            return chunks
            
        except Exception as e:
            text = document.text
            chunks = []
            
            for i in range(0, len(text), chunk_size - chunk_overlap):
                chunk_text = text[i:i + chunk_size]
                chunks.append(TextChunk(
                    text=chunk_text,
                    source=document.source
                ))
            
            return chunks

    def embed_chunks(self, chunks: List[TextChunk]) -> List[TextChunk]:
        texts = [chunk.text for chunk in chunks]
        embeddings = self.model.encode(texts)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
        
        return chunks

class RAGChatbot:
    def __init__(self):
        self.processor = DocumentProcessor()
        self.pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        self.index_name = "ragchatbot"
        self.current_namespace = None
        
        model_path = "MBZUAI/Llama-3-Nanda-10B-Chat"
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.tokenizer = AutoTokenizer.from_pretrained(model_path)
        self.model = AutoModelForCausalLM.from_pretrained(
            model_path, 
            device_map = "auto", 
            trust_remote_code = True,
            use_auth_token = os.getenv("USE_AUTH_TOKEN"),
            offload_folder = "/home/aditya/Documents/AgenticSystems101/offloader"
        )
        
        max_retries = 3
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                if self.index_name in self.pc.list_indexes().names():
                    self.pc.delete_index(self.index_name)
                    time.sleep(2)
                
                self.pc.create_index(
                    name=self.index_name,
                    dimension=768,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                
                while True:
                    index_desc = self.pc.describe_index(self.index_name)
                    if index_desc.status.get('ready', False):
                        break
                    time.sleep(1)
                    
                break
                
            except Exception as e:
                if attempt == max_retries - 1:
                    raise e
                time.sleep(retry_delay)
        
        self.index = self.pc.Index(self.index_name)

    def set_namespace(self, session_hash: str):
        self.current_namespace = f"session_{session_hash}"

    def cleanup_namespace(self):
        if self.current_namespace:
            try:
                stats = self.index.describe_index_stats()
                if self.current_namespace in stats.namespaces:
                    self.index.delete(delete_all=True, namespace=self.current_namespace)
            except Exception as e:
                if not (hasattr(e, 'code') and str(e.code) == '404'):
                    print(f"Error cleaning up namespace: {str(e)}")

    def process_uploaded_files(self, files, request: gr.Request) -> str:
        """Process uploaded files and store their embeddings in Pinecone."""
        if not request or not request.session_hash:
            return "Error: No session information available"
            
        self.set_namespace(request.session_hash)
        self.cleanup_namespace()
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                total_chunks = 0
                batch_size = 100
                vectors_batch = []
                
                for file in files:
                    file_path = os.path.join(temp_dir, os.path.basename(file.name))
                    shutil.copy(file.name, file_path)
                    
                    document = self.processor.read_file(file_path)
                    chunks = self.processor.chunk_text(document)
                    chunks = self.processor.embed_chunks(chunks)
                    
                    for idx, chunk in enumerate(chunks):
                        vector_id = f"chunk_{total_chunks + idx}"
                        vectors_batch.append(
                            (vector_id, 
                             chunk.embedding.tolist(),
                             {
                                 "text": chunk.text,
                                 "source": chunk.source
                             })
                        )
                        
                        if len(vectors_batch) >= batch_size:
                            self.index.upsert(
                                vectors=vectors_batch,
                                namespace=self.current_namespace
                            )
                            vectors_batch = []
                    
                    total_chunks += len(chunks)
                
                if vectors_batch:
                    self.index.upsert(
                        vectors=vectors_batch,
                        namespace=self.current_namespace
                    )
                
                max_wait = 10
                start_time = time.time()
                while time.time() - start_time < max_wait:
                    stats = self.index.describe_index_stats()
                    if stats.namespaces.get(self.current_namespace, {}).get('vector_count', 0) >= total_chunks:
                        break
                    time.sleep(1)

                return f"दस्तावेज़ों को सफलतापूर्वक प्रोसेस किया गया और {total_chunks} चंक्स को डेटाबेस में डाला गया। कृपया क्वेरी करने से पहले कुछ सेकंड प्रतीक्षा करें।"

        except Exception as e:
            return f"Error processing files: {str(e)}"

    def generate_response(self, prompt: str) -> str:
        input_ids = self.tokenizer(prompt, return_tensors="pt").input_ids
        inputs = input_ids.to(self.device)
        
        generate_ids = self.model.generate(
            inputs,
            top_p=0.95,
            temperature=0.2,
            max_new_tokens=768,
            repetition_penalty=1.3,
            do_sample=True,
            pad_token_id=self.tokenizer.pad_token_id,
        )
        
        response = self.tokenizer.batch_decode(
            generate_ids, 
            skip_special_tokens=True, 
            clean_up_tokenization_spaces=True
        )[0]
        
        if "Response:\n" in response:
            return response.split("Response:\n")[1]
        return response.split("Response:\n")[1]

    def chat(self, message: str, history: List[List[str]], request: gr.Request) -> Tuple[List[List[str]], str]:
        if not request or not request.session_hash:
            return history + [[message, "Error: No session information available"]], ""
            
        self.set_namespace(request.session_hash)
        
        try:
            query_embedding = self.processor.model.encode(message)
            search_results = self.index.query(
                vector=query_embedding.tolist(),
                top_k=8,
                include_metadata=True,
                namespace=self.current_namespace
            )
            
            if not search_results.matches:
                return history + [[message, "मैं अपलोड किए गए दस्तावेजों में कोई संबंधित जानकारी नहीं ढूंढ सका। कृपया सुनिश्चित करें कि दस्तावेज सही तरीके से प्रोसेस किए गए हैं और फिर से प्रयास करने से पहले कुछ सेकंड का इंतजार करें।"]], ""
            
            context_with_sources = []
            chunks_display = []
            for i, match in enumerate(search_results.matches, 1):
                source = os.path.basename(match.metadata['source'])
                context_with_sources.append(f"From {source}:\n{match.metadata['text']}")
                chunks_display.append(f"Chunk {i} (Score: {match.score:.3f}) from {source}:\n{match.metadata['text']}\n")
            
            context = "\n\n---\n\n".join(context_with_sources)
            chunks_text = "\n---\n".join(chunks_display)
            
            chat_history = ""
            for h in history:
                chat_history += f"User: {h[0]}\nAssistant: {h[1]}\n"
            
            prompt = f"<|begin_of_text|><|start_header_id|>Role: User<|end_header_id|><|start_header_id|>Instructions: You are a helpful assistant. Use the provided context to answer questions accurately. If the context doesn't contain enough information, say so and provide general knowledge while being clear about what comes from the documents.\n\nContext:\n{context}\n\nChat History:\n{chat_history}\nCurrent Question: {message}<|end_header_id|><|start_header_id|>Response:<|end_header_id|><|eot_id|>"
            
            answer = self.generate_response(prompt)
            
            return history + [[message, answer]], chunks_text

        except Exception as e:
            return history + [[message, f"Error generating response: {str(e)}"]], ""

def create_demo():
    chatbot = RAGChatbot()

    def user_input(message, history, request: gr.Request):
        if not message.strip(): 
            return "", history, ""
        return "", *chatbot.chat(message, history, request)

    def process_files(files, request: gr.Request):
        return chatbot.process_uploaded_files(files, request)

    with gr.Blocks(theme=gr.themes.Soft()) as demo:
        gr.Markdown("# नमस्ते!")
        gr.Markdown("अपनी पसंद का दस्तावेज़ यहां रखें और चैट करना शुरू करें")        

        with gr.Row():
            with gr.Column(scale=1):
                file_output = gr.File(
                    file_count="multiple",
                    label="दस्तावेज़ अपलोड करें",
                    file_types=[".txt", ".pdf", ".docx"]
                )
                upload_button = gr.Button("दस्तावेज़ प्रक्रिया करें")
                status_box = gr.Textbox(label="स्टेटस", interactive=False)

            with gr.Column(scale=2):
                chatbot_interface = gr.Chatbot(
                    label="चैट का इतिहास",
                    height=400,
                    bubble_full_width=False,
                )
                chunks_display = gr.Textbox(
                    label="दस्तावेज़ के टुकड़े पुनःप्राप्त",
                    interactive=False,
                    lines=10,
                )
                with gr.Row():
                    msg = gr.Textbox(
                        label="अपना मैसेज टाइप करें",
                        placeholder="अपलोड किए गए दस्तावेजों के बारे में मुझसे कुछ भी पूछें... (भेजने के लिए एंटर दबाएं)",
                        lines=1,
                        scale=4,
                        max_lines=1
                    )
                clear = gr.Button("क्लियर")

        msg.submit(
            fn=user_input,
            inputs=[msg, chatbot_interface],
            outputs=[msg, chatbot_interface, chunks_display],
        )
        
        upload_button.click(
            fn=process_files,
            inputs=[file_output],
            outputs=[status_box],
        )
        
        clear.click(
            lambda: (None, None), 
            None, 
            [chatbot_interface, chunks_display], 
            queue=False
        )

    return demo

if __name__ == "__main__":
    demo = create_demo()
    demo.launch(share=True)
